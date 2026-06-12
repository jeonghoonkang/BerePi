from __future__ import annotations

import csv
import hashlib
import html
import mimetypes
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".xml",
    ".html",
    ".htm",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".css",
    ".scss",
    ".sql",
    ".yaml",
    ".yml",
    ".ini",
    ".cfg",
    ".log",
}


@dataclass(frozen=True)
class ExtractedDocument:
    file_name: str
    file_path: str
    author: str
    content: str
    mime_type: str
    file_size: int
    sha256: str
    extracted_at: str


def extract_file(path: Path, author: str = "") -> ExtractedDocument:
    path = path.resolve()
    content = _extract_content(path)
    mime_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
    stat = path.stat()
    return ExtractedDocument(
        file_name=path.name,
        file_path=str(path),
        author=author.strip(),
        content=content.strip(),
        mime_type=mime_type,
        file_size=stat.st_size,
        sha256=_sha256(path),
        extracted_at=datetime.now().isoformat(timespec="seconds"),
    )


def write_markdown_record(document: ExtractedDocument, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = _safe_stem(Path(document.file_name).stem)
    target = output_dir / f"{stem}_{document.sha256[:12]}.md"
    body = [
        "---",
        f"file_name: {document.file_name}",
        f"file_path: {document.file_path}",
        f"author: {document.author}",
        f"mime_type: {document.mime_type}",
        f"file_size: {document.file_size}",
        f"sha256: {document.sha256}",
        f"extracted_at: {document.extracted_at}",
        "---",
        "",
        "# Extracted File Record",
        "",
        "## Summary",
        "",
        f"- File: `{document.file_name}`",
        f"- Author: `{document.author or 'unknown'}`",
        f"- Size: `{document.file_size}` bytes",
        f"- MIME: `{document.mime_type}`",
        "",
        "## Content",
        "",
        "```text",
        document.content[:200000] if document.content else "(no extractable text)",
        "```",
        "",
    ]
    target.write_text("\n".join(body), encoding="utf-8")
    return target


def _extract_content(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _read_text(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    return f"Binary file metadata only. No text extractor is configured for {suffix or 'this file type'}."


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _extract_xlsx(path: Path) -> str:
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return "Excel file detected. Install openpyxl to extract workbook text."

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    parts.append("\t".join(values))
    finally:
        workbook.close()
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError:
        return "DOCX file detected. Install python-docx to extract document text."

    document = docx.Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    tables = []
    for table in document.tables:
        for row in table.rows:
            tables.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs + tables)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            return "PDF file detected. Install pypdf to extract PDF text."

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"# Page {index}\n{text}")
    return "\n\n".join(pages)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_stem(value: str) -> str:
    normalized = re.sub(r"[^0-9A-Za-z가-힣_-]+", "_", value).strip("_")
    return normalized or "file"
